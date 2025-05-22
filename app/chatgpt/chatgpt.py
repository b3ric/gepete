from docx import Document
import os
from io import BytesIO
import openai


openai.organization = "org-5yY762ZV0aAg2rCUoD2sWQAq"
openai.api_key = os.getenv('OPENAI_API_KEY')


def chat_complete(prompt):
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "user", "content": f"{prompt}"},
        ],
        temperature=0.7,
        max_tokens=1000
    )

    return response

def img_gen(prompt):
    response = openai.Image.create(
        model="dall-e-3",
        prompt=f"{prompt}",
        n=1,
        size="1024x1024",
    )

    return response

def aud2text(aud_file, lang):
    response = openai.Audio.transcribe(
        file = aud_file,
        model='whisper-1',
        temperature = 0.7,
        language = lang
    )

    return response

def update_briefing(doc_file, prompt):
    doc = Document(doc_file)
    original_text = "\n".join([para.text for para in doc.paragraphs])

    full_prompt = f"""Update the following donor briefing with the notes provided below.

Original Briefing:
{original_text}

Meeting Notes:
{prompt}

Provide a polished, updated version of the briefing, ready to be shared.
"""

    response = openai.ChatCompletion.create(
        model="gpt-4",
        messages=[{"role": "user", "content": full_prompt}],
        temperature=0.5,
        max_tokens=1500,
    )

    updated_text = response['choices'][0]['message']['content']

    updated_doc = Document()
    for line in updated_text.split("\n"):
        updated_doc.add_paragraph(line.strip())

    buffer = BytesIO()
    updated_doc.save(buffer)
    buffer.seek(0)
    return buffer